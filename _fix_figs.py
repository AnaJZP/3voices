"""
Regenerate Figures 2, 3, 4 with layout/label fixes and insert into Word.
- Fig 2: Move stats boxes outside violin area, fix overlaps
- Fig 3: Fix first label overlap
- Fig 4: 2-row layout (top: academic + institutional, bottom: public centered)
"""
import sys; sys.stdout.reconfigure(encoding='utf-8', errors='replace')
import os; os.environ["CUDA_VISIBLE_DEVICES"] = ""

import numpy as np
import pandas as pd
import networkx as nx
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from constants import normalize_group_series

# ═══════════════════════════════════════════════════════════════
# FIGURE 2: VIOLIN — no overlapping labels
# ═══════════════════════════════════════════════════════════════
print("1. Figure 2 (violin)...")

sent = pd.read_csv("results/sentiment/sentiment_results.csv", encoding="utf-8-sig")
sent["source"] = normalize_group_series(sent["source"])
inst = pd.read_csv("data/raw/institutional/institutional_with_sentiment.csv",
                    encoding="utf-8-sig")
inst["source"] = "institutional"

combined = pd.concat([
    sent[sent["source"]=="academic"][["source","sentiment_score"]],
    inst[["source","sentiment_score"]],
    sent[sent["source"]=="public"][["source","sentiment_score"]],
], ignore_index=True).dropna(subset=["sentiment_score"])

plt.rcParams.update({"font.family": "sans-serif", "font.size": 12})

fig, ax = plt.subplots(figsize=(11, 7.5))

colors = {"academic": "#4A90D9", "institutional": "#F5A623", "public": "#7ED321"}
order = ["academic", "institutional", "public"]
labels_nice = ["Academic", "Institutional", "Public"]
data_lists = [combined[combined["source"]==s]["sentiment_score"].values for s in order]

# Violin
parts = ax.violinplot(data_lists, positions=[1,2,3], showmeans=False,
                      showextrema=False, widths=0.7)
for pc, src in zip(parts["bodies"], order):
    pc.set_facecolor(colors[src])
    pc.set_alpha(0.3)

# Box
bp = ax.boxplot(data_lists, positions=[1,2,3], widths=0.1,
                showfliers=False, patch_artist=True, zorder=3,
                medianprops=dict(color="white", linewidth=2.5),
                whiskerprops=dict(color="gray", linewidth=1.2),
                capprops=dict(color="gray", linewidth=1.2))
for patch, src in zip(bp["boxes"], order):
    patch.set_facecolor(colors[src])
    patch.set_alpha(0.85)

# Stats as a clean table-like annotation OUTSIDE the plot area (right side)
stats_text = ""
for i, src in enumerate(order):
    vals = combined[combined["source"]==src]["sentiment_score"]
    m, sd, n = vals.mean(), vals.std(), len(vals)
    stats_text += f"{labels_nice[i]:14s}  M={m:.3f}  SD={sd:.3f}  N={n:,}\n"

ax.text(1.02, 0.95, stats_text.strip(), transform=ax.transAxes,
        fontsize=10, fontfamily="monospace", va="top",
        bbox=dict(boxstyle="round,pad=0.4", facecolor="#F8F8F8",
                 edgecolor="#CCCCCC", alpha=0.95))

# Effect size brackets — compact, no overlap
# Academic ↔ Institutional
ax.annotate("", xy=(1, 1.06), xytext=(2, 1.06),
            arrowprops=dict(arrowstyle="-", color="#555", lw=1.2))
ax.plot([1,1], [1.04,1.06], color="#555", lw=1.2)
ax.plot([2,2], [1.04,1.06], color="#555", lw=1.2)
ax.text(1.5, 1.07, "\u03b4 = \u20130.61 (large)", ha="center", fontsize=10,
        color="#555", fontstyle="italic")

# Institutional ↔ Public
ax.annotate("", xy=(2, 1.13), xytext=(3, 1.13),
            arrowprops=dict(arrowstyle="-", color="#555", lw=1.2))
ax.plot([2,2], [1.11,1.13], color="#555", lw=1.2)
ax.plot([3,3], [1.11,1.13], color="#555", lw=1.2)
ax.text(2.5, 1.14, "\u03b4 = +0.63 (large)", ha="center", fontsize=10,
        color="#555", fontstyle="italic")

# Academic ↔ Public
ax.annotate("", xy=(1, 1.20), xytext=(3, 1.20),
            arrowprops=dict(arrowstyle="-", color="#999", lw=1.0, linestyle="--"))
ax.plot([1,1], [1.18,1.20], color="#999", lw=1.0)
ax.plot([3,3], [1.18,1.20], color="#999", lw=1.0)
ax.text(2.0, 1.21, "\u03b4 = +0.16 (small)", ha="center", fontsize=9.5,
        color="#999", fontstyle="italic")

ax.set_xticks([1,2,3])
ax.set_xticklabels(labels_nice, fontsize=14, fontweight="bold")
ax.set_ylabel("Sentiment Score [0, 1]", fontsize=14)
ax.set_ylim(-0.05, 1.30)
ax.set_xlim(0.3, 3.7)
ax.grid(axis="y", alpha=0.15, linestyle="--")
ax.spines["top"].set_visible(False)
ax.spines["right"].set_visible(False)

plt.tight_layout()
fig2_path = Path("figures/fig2_sentiment_violin_updated.png")
plt.savefig(fig2_path, dpi=300, bbox_inches="tight", facecolor="white")
plt.close()
print(f"   Saved: {fig2_path}")

# ═══════════════════════════════════════════════════════════════
# FIGURE 3: TEMPORAL — fix first label overlap
# ═══════════════════════════════════════════════════════════════
print("\n2. Figure 3 (temporal)...")

yearly = pd.read_csv("results/comparison/yearly_sentiment.csv", encoding="utf-8-sig")
acad = yearly[yearly["source"] == "academic"].sort_values("year")

fig, ax = plt.subplots(figsize=(14, 7))

years = acad["year"].values
means = acad["mean_sentiment"].values
stds  = acad["std_sentiment"].values

ax.fill_between(years, means - stds, means + stds,
                color="#4A90D9", alpha=0.15, label=r"$\pm$1 SD")
ax.plot(years, means, "o-", color="#4A90D9", linewidth=2.5,
        markersize=8, markerfacecolor="white", markeredgewidth=2,
        label="Mean sentiment", zorder=5)

# Events — alternate top/bottom, first label shifted right to avoid edge clip
events = [
    (2015, "SDGs",             "top"),
    (2016, "Paris\nAgreement", "bottom"),
    (2018, "IPCC 1.5\u00b0C",  "bottom"),
    (2019, "Green Deal",       "top"),
    (2020, "COVID-19",         "bottom"),
    (2021, "COP26",            "top"),
    (2022, "COP27",            "bottom"),
    (2023, "COP28",            "top"),
    (2024, "COP29",            "bottom"),
]

y_upper = max(means + stds) + 0.015
y_lower = min(means - stds) - 0.015

for yr, label, pos in events:
    ax.axvline(yr, color="#DDDDDD", linewidth=0.8, linestyle="--", zorder=1)
    
    if pos == "top":
        y = y_upper
        va = "bottom"
    else:
        y = y_lower
        va = "top"
    
    # Shift first label slightly right to avoid left-edge clipping
    x = yr + 0.15 if yr == 2015 else yr
    
    ax.text(x, y, label, ha="center", va=va,
            fontsize=10, color="#555", fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.25", facecolor="white",
                     edgecolor="#DDD", alpha=0.9))

# Trend
ax.annotate("", xy=(2025, means[-2]), xytext=(2015, means[0]),
            arrowprops=dict(arrowstyle="->", color="#E74C3C", lw=2))
drop_pct = (means[0] - means[-2]) / means[0] * 100
ax.text(2020, means[0] + 0.015,
        f"Decline: {drop_pct:.0f}% over the decade",
        fontsize=12, color="#E74C3C", fontweight="bold", ha="center")

ax.set_xlabel("Year", fontsize=14, fontweight="bold")
ax.set_ylabel("Mean Sentiment Score [0, 1]", fontsize=14, fontweight="bold")
ax.set_xlim(2014.5, 2026.5)
ax.set_ylim(y_lower - 0.06, y_upper + 0.08)
ax.xaxis.set_major_locator(mticker.MultipleLocator(1))
ax.tick_params(axis="both", labelsize=12)
ax.legend(loc="upper right", fontsize=12, framealpha=0.9)
ax.grid(axis="y", alpha=0.15, linestyle="--")
ax.spines["top"].set_visible(False)
ax.spines["right"].set_visible(False)

plt.tight_layout()
fig3_path = Path("figures/fig3_temporal_sentiment.png")
plt.savefig(fig3_path, dpi=300, bbox_inches="tight", facecolor="white")
plt.close()
print(f"   Saved: {fig3_path}")

# ═══════════════════════════════════════════════════════════════
# FIGURE 4: NETWORKS — 2 rows (top: acad+inst, bottom: public centered)
# ═══════════════════════════════════════════════════════════════
print("\n3. Figure 4 (networks, 2-row layout)...")

try:
    import community as community_louvain
except ImportError:
    community_louvain = None

G_acad = nx.read_graphml("results/networks/academic/network_academic.graphml")
G_pub  = nx.read_graphml("results/networks/public/network_public.graphml")
edges_inst = pd.read_csv("results/networks/institutional/edges_institutional.csv",
                          encoding="utf-8-sig")
G_inst = nx.Graph()
for _, row in edges_inst.iterrows():
    G_inst.add_edge(row["source"], row["target"], weight=row["weight"])

def get_top(G, n=50):
    d = dict(G.degree())
    top = sorted(d, key=d.get, reverse=True)[:n]
    return G.subgraph(top).copy()

PALETTE = ["#4A90D9","#F5A623","#7ED321","#E91E63","#9C27B0",
           "#00BCD4","#795548","#607D8B","#CDDC39","#FF5722"]

def draw_network(ax, G, title, top_n_labels=15):
    if community_louvain:
        part = community_louvain.best_partition(G, weight="weight", random_state=42)
    else:
        from networkx.algorithms.community import greedy_modularity_communities
        comms = list(greedy_modularity_communities(G, weight="weight"))
        part = {n: i for i, c in enumerate(comms) for n in c}
    
    node_colors = [PALETTE[part.get(n,0) % len(PALETTE)] for n in G.nodes()]
    
    w_deg = dict(G.degree(weight="weight"))
    max_wd = max(w_deg.values()) if w_deg else 1
    node_sizes = [250 + 3000 * (w_deg.get(n,0)/max_wd) for n in G.nodes()]
    
    ew = [float(d.get("weight",1)) for _,_,d in G.edges(data=True)]
    max_w = max(ew) if ew else 1
    edge_w = [0.1 + 2.0*(w/max_w) for w in ew]
    ea = 0.03 if "0.99" in title else 0.10
    
    k = 5.5 / np.sqrt(max(G.number_of_nodes(),1))
    pos = nx.spring_layout(G, k=k, iterations=300, seed=42, weight="weight")
    
    nx.draw_networkx_edges(G, pos, ax=ax, alpha=ea, width=edge_w,
                           edge_color="#BBBBBB")
    nx.draw_networkx_nodes(G, pos, ax=ax, node_color=node_colors,
                           node_size=node_sizes, alpha=0.85,
                           linewidths=0.8, edgecolors="white")
    
    top_nodes = sorted(w_deg, key=w_deg.get, reverse=True)[:top_n_labels]
    for node in top_nodes:
        x, y = pos[node]
        ax.text(x, y + 0.035, node, fontsize=12, fontweight="bold",
                ha="center", va="bottom", color="#222",
                bbox=dict(boxstyle="round,pad=0.2", facecolor="white",
                         edgecolor="none", alpha=0.85))
    
    ax.set_title(title, fontsize=14, fontweight="bold", pad=12)
    ax.axis("off")

# Layout: 2 rows. Row 1 = 2 panels, Row 2 = 1 panel centered
fig = plt.figure(figsize=(18, 18))

# Top left: Academic
ax1 = fig.add_subplot(2, 2, 1)
draw_network(ax1, get_top(G_acad),
             "Academic  |  Q = 0.42  |  \u03c1 = 0.63  |  5 comm.")

# Top right: Institutional
ax2 = fig.add_subplot(2, 2, 2)
draw_network(ax2, get_top(G_inst),
             "Institutional  |  Q = 0.08  |  \u03c1 = 0.99  |  4 comm.")

# Bottom center: Public (spans middle of bottom row)
ax3 = fig.add_axes([0.25, 0.02, 0.50, 0.45])  # [left, bottom, width, height]
draw_network(ax3, get_top(G_pub),
             "Public  |  Q = 0.11  |  \u03c1 = 0.58  |  4 comm.")

plt.suptitle("Semantic Co-occurrence Networks (Top 50 Nodes by Degree)",
             fontsize=16, fontweight="bold", y=0.98)

fig4_path = Path("figures/fig4_semantic_networks_three.png")
plt.savefig(fig4_path, dpi=300, bbox_inches="tight", facecolor="white")
plt.close()
print(f"   Saved: {fig4_path}")

# ═══════════════════════════════════════════════════════════════
# INSERT INTO WORD
# ═══════════════════════════════════════════════════════════════
print("\n4. Inserting into Word...")
doc = Document(DOC_PATH)

def pidx(text, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start: continue
        if text.lower() in p.text.lower():
            return i
    return None

def replace_fig(caption_text, fig_path, width, start=0):
    cap = pidx(caption_text, start)
    if cap and cap > 0:
        prev = doc.paragraphs[cap - 1]
        for run in prev.runs: run.clear()
        for d in prev._element.findall('.//' + qn('w:drawing')):
            d.getparent().remove(d)
        run = prev.add_run()
        run.add_picture(str(fig_path), width=width)
        print(f"   Replaced {caption_text} at paragraph {cap-1}")
        return True
    print(f"   WARNING: {caption_text} not found")
    return False

replace_fig("Figure 2", fig2_path, Inches(5.5), start=80)
replace_fig("Figure 3", fig3_path, Inches(6.0), start=88)
replace_fig("Fig. 4",   fig4_path, Inches(6.0), start=95)

doc.save(DOC_PATH)
print(f"\n   Saved: {DOC_PATH}")
print("   Done!")
